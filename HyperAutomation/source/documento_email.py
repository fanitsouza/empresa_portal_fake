from docx import Document
import smtplib
from email.message import EmailMessage
from dotenv import load_dotenv
import os
from pathlib import Path

# Localiza o arquivo .env na pasta do projeto
BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"
load_dotenv(ENV_PATH)

def criar_documento(dados=None):
    documento = Document()
    documento.add_heading("PORTAL FAKE SOLUÇÕES DIGITAIS", level=1)
    documento.add_heading("FICHA DE CADASTRO", level=2)
    documento.add_paragraph()

    if dados:
        documento.add_paragraph(f"1. Nome: {dados.get('Nome', '')}")
        documento.add_paragraph(f"2. Sobrenome: {dados.get('Sobrenome', '')}")
        documento.add_paragraph(f"3. CPF: {dados.get('CPF', '')}")
        documento.add_paragraph(f"4. E-mail: {dados.get('E-mail', '')}")
        documento.add_paragraph(f"5. Telefone: {dados.get('Telefone', '')}")
        documento.add_paragraph(f"6. Data de Nascimento: {dados.get('Nascimento', '')}")
        documento.add_paragraph(f"7. Endereço: {dados.get('Endereco', '')}")
        cpf_limpo = (dados.get('CPF') or 'temp').replace('.', '').replace('-', '')
        nome_arquivo = f"Ficha_Cadastro_{cpf_limpo}.docx"
    else:
        documento.add_paragraph("1. Nome:")
        documento.add_paragraph("____________________________________________________________")
        documento.add_paragraph("2. Sobrenome:")
        documento.add_paragraph("____________________________________________________________")
        documento.add_paragraph("3. CPF:")
        documento.add_paragraph("____________________________________________________________")
        documento.add_paragraph("4. E-mail:")
        documento.add_paragraph("____________________________________________________________")
        documento.add_paragraph("5. Telefone:")
        documento.add_paragraph("____________________________________________________________")
        documento.add_paragraph("6. Data de Nascimento:")
        documento.add_paragraph("____________________________________________________________")
        documento.add_paragraph("7. Endereço:")
        documento.add_paragraph("____________________________________________________________")
        nome_arquivo = "Ficha_Cadastro_Portal_Fake.docx"

    documento.add_paragraph()
    documento.add_paragraph("Assinatura:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph()
    documento.add_paragraph("Data:")
    documento.add_paragraph("______/______/________")

    arquivo = BASE_DIR / nome_arquivo
    documento.save(str(arquivo))
    return str(arquivo)

def enviar_email(email_cliente, arquivo, apagar_apos_envio=True):
    remetente = os.getenv("EMAIL_REMETENTE")
    senha = os.getenv("EMAIL_SENHA")

    if not remetente:
        raise Exception("EMAIL_REMETENTE não encontrado no arquivo .env")
    if not senha:
        raise Exception("EMAIL_SENHA não encontrada no arquivo .env")

    print(f"Remetente: {remetente}")
    print(f"Destinatário: {email_cliente}")

    mensagem = EmailMessage()
    mensagem["Subject"] = "Ficha de Cadastro - Portal Fake Soluções Digitais"
    mensagem["From"] = remetente
    mensagem["To"] = email_cliente
    mensagem.set_content(
        "Prezado(a),\n\n"
        "Segue em anexo a ficha de cadastro da Empresa Portal Fake Soluções Digitais.\n"
        "Solicitamos que todos os campos sejam conferidos e devolvidos juntamente com:\n"
        "• Documento oficial com foto;\n"
        "• Comprovante de residência atualizado.\n\n"
        "Em caso de dúvidas, estamos à disposição.\n\n"
        "Atenciosamente,\n"
        "Portal Fake Soluções Digitais"
    )

    try:
        with open(arquivo, "rb") as doc:
            mensagem.add_attachment(
                doc.read(),
                maintype="application",
                subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=Path(arquivo).name
            )

        with smtplib.SMTP("smtp.gmail.com", 587) as servidor:
            servidor.starttls()
            servidor.login(remetente, senha)
            servidor.send_message(mensagem)

        print("E-mail enviado com sucesso!")
    finally:
        if apagar_apos_envio and os.path.exists(arquivo):
            os.remove(arquivo)
            print(f"🗑️ Arquivo temporário '{Path(arquivo).name}' removido após o envio.")

if __name__ == "__main__":
    dados_exemplo = {
        "Nome": "Ana",
        "Sobrenome": "Silva",
        "CPF": "10000012482",
        "E-mail": "ana.silva.1@exemplo.com",
        "Telefone": "(92) 99001-1001",
        "Nascimento": "1988-01-01",
        "Endereco": "Rua A, 100 - Bairro Centro - Manaus/AM"
    }
    arq = criar_documento(dados_exemplo)
    enviar_email("carvalhosannyer@gmail.com", arq, apagar_apos_envio=True)
